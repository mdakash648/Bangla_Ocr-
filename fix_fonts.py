"""
fix_fonts.py

Overview
- Purpose: Ensure true English words are formatted with Times New Roman, while all other content (including Bijoy/SutonnyMJ-style Bangla using ASCII) remains in SutonnyMJ, without disturbing the layout.

Approach
- Load the DOCX and walk paragraphs in the document body, tables (recursively), and headers/footers.
- For each run, split its text into chunks using three token types to preserve layout:
  1) English-like word: [A-Za-z][A-Za-z'-]+
  2) Whitespace: \\s+
  3) Anything else: [^A-Za-z\s]+
- For each chunk, create a new run immediately after the original run to maintain order, cloning style attributes. Set font to Times New Roman if the chunk is a true English word according to wordfreq, else set SutonnyMJ.

Safety and Limitations
- Original file is never overwritten; output is written to a new file.
- Runs inside hyperlinks and field codes are skipped to avoid corrupting link or field structure.
- python-docx does not expose text boxes/shapes (drawing layer) content for editing; those are skipped gracefully.
- Mixed-font documents are handled: only the newly created runs are explicitly given Times New Roman or SutonnyMJ; other formatting is preserved as much as python-docx allows.
- Control characters are sanitized from run text before processing.

Dependencies
- python-docx, regex, wordfreq, argparse, logging
"""

from __future__ import annotations

import argparse
import logging
import os
import sys
import time
from typing import List, Tuple
from copy import deepcopy

import regex as re
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
try:
    # Make wordfreq optional so GUI sentence mode can run without it
    from wordfreq import zipf_frequency as _zipf_frequency
except Exception:
    _zipf_frequency = None
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
try:
    # GUI is optional; only used when --gui is provided
    import tkinter as tk
    from tkinter import filedialog, messagebox, scrolledtext
except Exception:
    tk = None  # Graceful fallback for environments without tkinter


# Compile tokenization patterns once
EN_LIKE_WORD_RE = re.compile(r"[A-Za-z][A-Za-z'-]+")
TOKEN_SPLIT_RE = re.compile(r"([A-Za-z][A-Za-z'-]+|\s+|[^A-Za-z\s]+)")
CTRL_CHARS_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]")


def is_english_word(token: str, threshold: float) -> bool:
    """Return True if token is a true English word by frequency threshold.

    - Keep only [A-Za-z'-].
    - Length must be >= 2.
    - Use wordfreq.zipf_frequency(token.lower(), 'en') >= threshold.
    """
    if not token:
        return False
    cleaned = re.sub(r"[^A-Za-z'-]", "", token)
    if len(cleaned) < 2:
        return False
    if _zipf_frequency is None:
        # Signal that English detection is unavailable (handled by CLI gate)
        raise RuntimeError("wordfreq not installed; English detection unavailable.")
    freq = _zipf_frequency(cleaned.lower(), "en")
    return freq >= threshold


def _clean_token(token: str) -> str:
    return re.sub(r"[^A-Za-z'-]", "", token) if token else ""


def should_use_tnr(token: str, threshold: float, src_font_name: str | None) -> bool:
    """Decide whether to set Times New Roman for this token.

    Heuristic to avoid false positives for Bijoy-encoded Bangla:
    - Use is_english_word() as primary signal.
    - If the source run is explicitly SutonnyMJ and the cleaned token length <= 3,
      do NOT convert to TNR even if frequency passes threshold. This prevents
      converting short Bijoy syllables like "am", "to", "in", etc.
    - For length >= 4, follow is_english_word() result.
    """
    cleaned = _clean_token(token)
    if len(cleaned) < 2:
        return False

    if not is_english_word(cleaned, threshold):
        return False

    font_name = (src_font_name or "").lower()
    if "sutonny" in font_name and len(cleaned) <= 3:
        return False

    return True


def split_into_chunks(text: str) -> List[str]:
    """Split text into chunks preserving order and token types.

    Chunks follow this precedence order:
    - English-like word: [A-Za-z][A-Za-z'-]+
    - Whitespace: \\s+
    - Anything else: [^A-Za-z\s]+

    Returns a list of non-empty chunks.
    """
    if not text:
        return []
    parts = TOKEN_SPLIT_RE.findall(text)
    return [p for p in parts if p]


def sanitize_text(text: str) -> str:
    """Remove non-ASCII control characters except standard whitespace."""
    if not text:
        return ""
    return CTRL_CHARS_RE.sub("", text)


def clone_run_style(src_run: Run, dst_run: Run) -> None:
    """Clone key style attributes from src_run to dst_run.

    Copies: bold, italic, underline, size, color, all-caps, small-caps, style.
    Does not copy font.name (we set explicitly to Times New Roman or SutonnyMJ).
    """
    try:
        dst_run.style = src_run.style
    except Exception:
        # Style might not be applicable across parts; ignore quietly
        pass

    s_font = src_run.font
    d_font = dst_run.font

    d_font.bold = s_font.bold
    d_font.italic = s_font.italic
    d_font.underline = s_font.underline
    d_font.size = s_font.size
    d_font.all_caps = s_font.all_caps
    d_font.small_caps = s_font.small_caps

    # Copy color (either rgb or theme). If neither is present, leave default.
    try:
        if s_font.color is not None:
            if getattr(s_font.color, "rgb", None) is not None:
                d_font.color.rgb = s_font.color.rgb
            if getattr(s_font.color, "theme_color", None) is not None:
                d_font.color.theme_color = s_font.color.theme_color
    except Exception:
        pass

    # Preserve original rFonts mapping (ascii/hAnsi/eastAsia/cs/hint)
    copy_rFonts(src_run, dst_run)


def copy_rFonts(src_run: Run, dst_run: Run) -> None:
    """Copy the full rFonts element from src_run to dst_run.

    This preserves ascii, hAnsi, eastAsia, cs, and hint exactly as present
    in the source run, avoiding layout corruption for Bijoy/SutonnyMJ text.
    """
    try:
        s_rPr = getattr(src_run._r, "rPr", None)
        if s_rPr is None or getattr(s_rPr, "rFonts", None) is None:
            return
        src_rFonts = s_rPr.rFonts

        d_rPr = dst_run._r.get_or_add_rPr()
        # Remove existing rFonts if any
        if getattr(d_rPr, "rFonts", None) is not None:
            d_rPr.remove(d_rPr.rFonts)

        cloned = deepcopy(src_rFonts)
        d_rPr.append(cloned)
    except Exception:
        # Be fail-safe; if copy fails, silently continue
        pass


def set_ascii_hAnsi(run: Run, font_name: str) -> None:
    """Set only ascii and hAnsi of rFonts to font_name, leaving eastAsia/cs intact."""
    try:
        rPr = run._r.get_or_add_rPr()
        rFonts = getattr(rPr, "rFonts", None)
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)

        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)

        # Avoid theme overrides that may conflict
        if rFonts.get(qn("w:asciiTheme")) is not None:
            rFonts.attrib.pop(qn("w:asciiTheme"), None)
        if rFonts.get(qn("w:hAnsiTheme")) is not None:
            rFonts.attrib.pop(qn("w:hAnsiTheme"), None)
    except Exception:
        pass


def _is_run_in_hyperlink(run: Run) -> bool:
    """Detect if run is inside a hyperlink element to avoid breaking links."""
    try:
        parent = run._r.getparent()
        if parent is None:
            return False
        tag = getattr(parent, "tag", "")
        return tag.endswith("}hyperlink")
    except Exception:
        return False


def _is_run_field_code(run: Run) -> bool:
    """Detect field code runs (fldChar/instrText/fldSimple) and skip them."""
    r = run._r
    try:
        # Any fldChar or instrText descendants
        if r.xpath(".//w:fldChar", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
            return True
        if r.xpath(".//w:instrText", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
            return True
        # Parent fldSimple
        parent = r.getparent()
        if parent is not None:
            tag = getattr(parent, "tag", "")
            if tag.endswith("}fldSimple"):
                return True
    except Exception:
        return False
    return False


def _insert_run_after(paragraph: Paragraph, anchor_r) -> Run:
    """Create a new run and insert it immediately after anchor_r in the XML tree.

    Returns the new Run object.
    """
    new_run = paragraph.add_run("")  # creates at end, will reposition
    try:
        anchor_r.addnext(new_run._r)
    except Exception:
        # Fallback: leave at end if reposition fails
        pass
    return new_run


def process_run(paragraph: Paragraph, run: Run, threshold: float) -> Tuple[int, int]:
    """Process a single run by splitting and recreating styled runs.

    Returns a tuple: (runs_scanned, english_tokens_changed)
    """
    text = sanitize_text(run.text)
    if text == "":
        return (0, 0)

    if _is_run_in_hyperlink(run):
        logging.debug("Skipping run inside hyperlink: %r", text)
        return (1, 0)

    if _is_run_field_code(run):
        logging.debug("Skipping field code run: %r", text)
        return (1, 0)

    chunks = split_into_chunks(text)
    if not chunks:
        return (1, 0)

    english_changed = 0
    anchor = run._r

    src_font_name = None
    try:
        src_font_name = (run.font.name or "")
    except Exception:
        src_font_name = None

    for chunk in chunks:
        is_en_like = bool(EN_LIKE_WORD_RE.fullmatch(chunk))
        make_tnr = is_en_like and should_use_tnr(chunk, threshold, src_font_name)

        new_run = _insert_run_after(paragraph, anchor)
        new_run.text = chunk
        clone_run_style(run, new_run)

        try:
            if make_tnr:
                set_ascii_hAnsi(new_run, "Times New Roman")
                english_changed += 1
                logging.debug("Chunk ascii/hAnsi -> Times New Roman: %r", chunk)
            else:
                # Preserve original rFonts (already copied in clone_run_style)
                logging.debug("Chunk preserves original rFonts: %r", chunk)
        except Exception as e:
            logging.debug("Failed to adjust rFonts for chunk %r: %s", chunk, e)

        anchor = new_run._r

    # Clear original run text so we don't duplicate content
    run.text = ""

    return (1, english_changed)


def process_paragraph(paragraph: Paragraph, threshold: float) -> Tuple[int, int]:
    """Process all runs in a paragraph.

    Returns: (runs_scanned, english_tokens_changed)
    """
    runs_scanned = 0
    english_changed = 0

    # Work on a copy of the run list since we modify the paragraph
    original_runs = list(paragraph.runs)

    for run in original_runs:
        r_scanned, en_changed = process_run(paragraph, run, threshold)
        runs_scanned += r_scanned
        english_changed += en_changed

    return runs_scanned, english_changed


def process_table(table, threshold: float) -> Tuple[int, int, int]:
    """Process all paragraphs in a table (recursively includes nested tables).

    Returns: (paragraphs_processed, runs_scanned, english_tokens_changed)
    """
    paragraphs = 0
    runs_scanned = 0
    english_changed = 0

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                paragraphs += 1
                r_scanned, en_changed = process_paragraph(p, threshold)
                runs_scanned += r_scanned
                english_changed += en_changed
            for nested in cell.tables:
                p2, r2, e2 = process_table(nested, threshold)
                paragraphs += p2
                runs_scanned += r2
                english_changed += e2

    return paragraphs, runs_scanned, english_changed


def process_headers_footers(doc: Document, threshold: float) -> Tuple[Tuple[int, int, int], Tuple[int, int, int]]:
    """Process headers and footers of all sections.

    Returns ((hdr_paras, hdr_runs, hdr_english), (ftr_paras, ftr_runs, ftr_english))
    """
    hdr_paras = hdr_runs = hdr_en = 0
    ftr_paras = ftr_runs = ftr_en = 0

    for section in doc.sections:
        header = section.header
        if header is not None:
            for p in header.paragraphs:
                hdr_paras += 1
                r, e = process_paragraph(p, threshold)
                hdr_runs += r
                hdr_en += e
            for t in header.tables:
                p2, r2, e2 = process_table(t, threshold)
                hdr_paras += p2
                hdr_runs += r2
                hdr_en += e2

        footer = section.footer
        if footer is not None:
            for p in footer.paragraphs:
                ftr_paras += 1
                r, e = process_paragraph(p, threshold)
                ftr_runs += r
                ftr_en += e
            for t in footer.tables:
                p2, r2, e2 = process_table(t, threshold)
                ftr_paras += p2
                ftr_runs += r2
                ftr_en += e2

    return (hdr_paras, hdr_runs, hdr_en), (ftr_paras, ftr_runs, ftr_en)


def convert_docx(in_path: str, out_path: str | None, threshold: float, dry_run: bool) -> None:
    start = time.time()

    if not os.path.isfile(in_path):
        logging.error("Input file not found: %s", in_path)
        raise FileNotFoundError(f"Input file not found: {in_path}")

    try:
        doc = Document(in_path)
    except Exception as e:
        logging.error("Failed to open DOCX: %s", e)
        raise

    total_paras = 0
    total_runs = 0
    total_english_tokens = 0

    # Body paragraphs
    for p in doc.paragraphs:
        total_paras += 1
        r_scanned, en_changed = process_paragraph(p, threshold)
        total_runs += r_scanned
        total_english_tokens += en_changed

    # Tables in body
    for t in doc.tables:
        p2, r2, e2 = process_table(t, threshold)
        total_paras += p2
        total_runs += r2
        total_english_tokens += e2

    # Headers and footers
    (hdr_p, hdr_r, hdr_e), (ftr_p, ftr_r, ftr_e) = process_headers_footers(doc, threshold)

    # Log per-section counts for dry-run insight
    logging.info("Main body: paragraphs=%d, runs=%d, english_tokens=%d", total_paras, total_runs, total_english_tokens)
    logging.info("Headers: paragraphs=%d, runs=%d, english_tokens=%d", hdr_p, hdr_r, hdr_e)
    logging.info("Footers: paragraphs=%d, runs=%d, english_tokens=%d", ftr_p, ftr_r, ftr_e)

    grand_paras = total_paras + hdr_p + ftr_p
    grand_runs = total_runs + hdr_r + ftr_r
    grand_english = total_english_tokens + hdr_e + ftr_e

    elapsed = time.time() - start

    if dry_run:
        logging.info("Dry-run: no file written. Total paragraphs=%d, runs=%d, english tokens changed=%d, elapsed=%.2fs",
                     grand_paras, grand_runs, grand_english, elapsed)
        return

    # Determine output path
    if not out_path:
        base, ext = os.path.splitext(in_path)
        out_path = f"{base}_fixed{ext or '.docx'}"

    # Safety: never overwrite input
    if os.path.abspath(out_path) == os.path.abspath(in_path):
        base, ext = os.path.splitext(in_path)
        out_path = f"{base}_fixed{ext or '.docx'}"
        logging.warning("Output path same as input; writing to %s", out_path)

    try:
        doc.save(out_path)
    except Exception as e:
        logging.error("Failed to save output DOCX: %s", e)
        raise

    logging.info(
        "Wrote: %s | Total paragraphs=%d, runs=%d, english tokens changed=%d, elapsed=%.2fs",
        out_path, grand_paras, grand_runs, grand_english, elapsed,
    )


# =========================
# Sentence-based processing
# =========================

SENTENCE_SPLIT_RE = re.compile(r"(?<!\w\.[A-Za-z])(?<![A-Z][a-z]\.)\s*(.+?[.!?])(?:\s+|$)")


def extract_sentences(text: str) -> List[str]:
    """Extract complete sentences (ending with .!?) from input text.

    Returns list of trimmed sentences including their terminal punctuation.
    """
    if not text:
        return []
    sentences = [m.group(1).strip() for m in SENTENCE_SPLIT_RE.finditer(text)]
    # Deduplicate while preserving order
    seen = set()
    ordered: List[str] = []
    for s in sentences:
        if s and s not in seen:
            ordered.append(s)
            seen.add(s)
    return ordered


def build_paragraph_text_map(paragraph: Paragraph) -> Tuple[str, List[Tuple[Run, int, int]]]:
    """Return (full_text, mapping) for a paragraph.

    mapping is a list of tuples: (run_object, start_index_in_full_text, end_index).
    """
    full = []
    mapping: List[Tuple[Run, int, int]] = []
    cursor = 0
    for r in paragraph.runs:
        t = r.text or ""
        full.append(t)
        start = cursor
        end = cursor + len(t)
        mapping.append((r, start, end))
        cursor = end
    return ("".join(full), mapping)


def split_run_apply_tnr(paragraph: Paragraph, run: Run, local_start: int, local_end: int) -> None:
    """Split a run at [local_start:local_end) and set TNR for that slice.

    The original run is cleared and replaced with up to three new runs to
    preserve formatting around the targeted slice.
    """
    # Preserve document structure for hyperlinks/field codes
    if _is_run_in_hyperlink(run) or _is_run_field_code(run):
        return

    text = run.text or ""
    if local_start < 0:
        local_start = 0
    if local_end > len(text):
        local_end = len(text)
    if local_start >= local_end:
        return

    before = text[:local_start]
    target = text[local_start:local_end]
    after = text[local_end:]

    anchor = run._r

    if before:
        r_before = _insert_run_after(paragraph, anchor)
        r_before.text = before
        clone_run_style(run, r_before)
        anchor = r_before._r

    r_target = _insert_run_after(paragraph, anchor)
    r_target.text = target
    clone_run_style(run, r_target)
    set_ascii_hAnsi(r_target, "Times New Roman")
    anchor = r_target._r

    if after:
        r_after = _insert_run_after(paragraph, anchor)
        r_after.text = after
        clone_run_style(run, r_after)

    # Clear original run text
    run.text = ""


def apply_sentence_to_paragraph(paragraph: Paragraph, sentence: str) -> int:
    """Find and convert all occurrences of `sentence` in this paragraph.

    Returns the number of matches converted.
    """
    if not sentence:
        return 0

    converted = 0
    search_from = 0

    while True:
        full_text, mapping = build_paragraph_text_map(paragraph)
        if not full_text or search_from >= len(full_text):
            break

        # Escape for literal match; we want exact sentence matches
        pattern = re.compile(re.escape(sentence))
        m = pattern.search(full_text, pos=search_from)
        if not m:
            break

        g_start, g_end = m.span()

        # Apply to all overlapping runs
        for run, r_start, r_end in mapping:
            overlap_start = max(g_start, r_start)
            overlap_end = min(g_end, r_end)
            if overlap_start < overlap_end:
                local_start = overlap_start - r_start
                local_end = overlap_end - r_start
                split_run_apply_tnr(paragraph, run, local_start, local_end)

        converted += 1
        # Move search_from past this match in the updated paragraph
        search_from = g_end

    return converted


def apply_sentences_docx(in_path: str, out_path: str | None, sentences: List[str]) -> Tuple[str, int]:
    """Apply Times New Roman to each provided sentence across the document.

    Returns (written_path, total_converted_matches).
    """
    if not os.path.isfile(in_path):
        raise FileNotFoundError(f"Input file not found: {in_path}")

    doc = Document(in_path)
    total_converted = 0

    # Helper to process all paragraphs recursively
    def process_paragraphs(paragraphs: List[Paragraph]) -> None:
        nonlocal total_converted
        for p in paragraphs:
            for s in sentences:
                total_converted += apply_sentence_to_paragraph(p, s)

    # Body
    process_paragraphs(doc.paragraphs)
    # Tables in body
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)
                for nt in cell.tables:
                    for nrow in nt.rows:
                        for ncell in nrow.cells:
                            process_paragraphs(ncell.paragraphs)

    # Headers/Footers
    for section in doc.sections:
        hdr = section.header
        if hdr is not None:
            process_paragraphs(hdr.paragraphs)
            for t in hdr.tables:
                for row in t.rows:
                    for cell in row.cells:
                        process_paragraphs(cell.paragraphs)
        ftr = section.footer
        if ftr is not None:
            process_paragraphs(ftr.paragraphs)
            for t in ftr.tables:
                for row in t.rows:
                    for cell in row.cells:
                        process_paragraphs(cell.paragraphs)

    # Output path
    base, ext = os.path.splitext(in_path)
    if not out_path:
        out_path = f"{base}_fixed{ext or '.docx'}"
    if os.path.abspath(out_path) == os.path.abspath(in_path):
        out_path = f"{base}_fixed{ext or '.docx'}"

    doc.save(out_path)
    return out_path, total_converted


# =======
#   GUI
# =======

def launch_gui() -> None:
    if tk is None:
        print("tkinter is not available in this environment.")
        sys.exit(3)

    root = tk.Tk()
    root.title("Fix Fonts - Sentence Mode")

    # File selection
    file_frame = tk.Frame(root)
    file_frame.pack(fill="x", padx=8, pady=6)

    tk.Label(file_frame, text="Input .docx:").pack(side="left")
    in_var = tk.StringVar(value="")
    in_entry = tk.Entry(file_frame, textvariable=in_var, width=60)
    in_entry.pack(side="left", padx=6)

    def browse_in():
        path = filedialog.askopenfilename(title="Select DOCX", filetypes=[["Word Document", "*.docx"]])
        if path:
            in_var.set(path)

    tk.Button(file_frame, text="Browse", command=browse_in).pack(side="left")

    # Output path (optional)
    out_frame = tk.Frame(root)
    out_frame.pack(fill="x", padx=8, pady=0)
    tk.Label(out_frame, text="Output .docx (optional):").pack(side="left")
    out_var = tk.StringVar(value="")
    out_entry = tk.Entry(out_frame, textvariable=out_var, width=60)
    out_entry.pack(side="left", padx=6)

    def browse_out():
        path = filedialog.asksaveasfilename(title="Save As", defaultextension=".docx", filetypes=[["Word Document", "*.docx"]])
        if path:
            out_var.set(path)

    tk.Button(out_frame, text="Save As", command=browse_out).pack(side="left")

    # Text area for sentences
    tk.Label(root, text="Enter English sentences (free text or one per line):").pack(anchor="w", padx=8, pady=(8, 0))
    txt = scrolledtext.ScrolledText(root, width=90, height=12, wrap="word")
    txt.pack(fill="both", expand=True, padx=8, pady=6)

    # Status
    status_var = tk.StringVar(value="Ready")
    status = tk.Label(root, textvariable=status_var, anchor="w")
    status.pack(fill="x", padx=8, pady=(0, 6))

    # Actions
    btn_frame = tk.Frame(root)
    btn_frame.pack(fill="x", padx=8, pady=(0, 10))

    def do_split():
        sentences = extract_sentences(txt.get("1.0", "end").strip())
        txt.delete("1.0", "end")
        txt.insert("1.0", "\n".join(sentences))
        status_var.set(f"Separated {len(sentences)} sentence(s)")

    def do_clear():
        txt.delete("1.0", "end")
        status_var.set("Cleared input")

    def do_convert():
        in_path = in_var.get().strip()
        out_path = out_var.get().strip() or None
        if not in_path:
            messagebox.showerror("Missing file", "Please choose an input .docx file.")
            return
        raw = txt.get("1.0", "end").strip()
        # Accept either one sentence per line or free text; split robustly
        sentences = [s for s in (line.strip() for line in raw.splitlines()) if s]
        # If text wasn't line-separated, extract by regex
        if len(sentences) <= 1:
            sentences = extract_sentences(raw)
        if not sentences:
            messagebox.showwarning("No sentences", "Provide at least one sentence ending with .!? ")
            return
        try:
            status_var.set("Processing...")
            root.update_idletasks()
            written, converted = apply_sentences_docx(in_path, out_path, sentences)
            status_var.set(f"Done. Converted {converted} match(es). Wrote: {written}")
            messagebox.showinfo("Success", f"Converted {converted} match(es).\nSaved: {written}")
        except Exception as e:
            messagebox.showerror("Error", str(e))
            status_var.set("Error: see popup")

    tk.Button(btn_frame, text="Separate Sentences", command=do_split).pack(side="left")
    tk.Button(btn_frame, text="Convert Fonts", command=do_convert).pack(side="left", padx=8)
    tk.Button(btn_frame, text="Clear", command=do_clear).pack(side="left")

    root.mainloop()


def build_arg_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(description="Fix fonts in DOCX: English -> Times New Roman; other -> SutonnyMJ")
    p.add_argument("--in", dest="in_path", required=False, help="Input .docx file path")
    p.add_argument("--out", dest="out_path", default=None, help="Output .docx file path (optional)")
    p.add_argument("--threshold", dest="threshold", type=float, default=3.5,
                   help="Zipf frequency threshold for English detection (default: 3.5)")
    p.add_argument("--dry-run", dest="dry_run", action="store_true", help="Analyze only; do not write output")
    p.add_argument("--verbose", dest="verbose", action="store_true", help="Enable DEBUG logging")
    p.add_argument("--gui", dest="gui", action="store_true", help="Launch GUI for sentence-based conversion")
    p.add_argument("--sentences", dest="sentences", default=None,
                   help="Comma-separated English sentences to convert (exact matches)")
    p.add_argument("--sentences-file", dest="sentences_file", default=None,
                   help="Path to a text file containing sentences (one per line or free text)")
    return p


def main(argv: List[str] | None = None) -> None:
    parser = build_arg_parser()
    args = parser.parse_args(argv)

    level = logging.DEBUG if args.verbose else logging.INFO
    logging.basicConfig(level=level, format="%(levelname)s: %(message)s")

    logging.debug("Arguments: %s", args)

    if getattr(args, "gui", False):
        launch_gui()
        return

    # Sentence-based CLI mode
    if args.sentences or args.sentences_file:
        if not args.in_path:
            logging.error("--in is required when using --sentences or --sentences-file")
            sys.exit(2)
        sentences: List[str] = []
        if args.sentences:
            # Split on commas but allow commas inside quotes is overkill; keep simple
            raw_items = [s.strip() for s in args.sentences.split(",")]
            for item in raw_items:
                if item:
                    # If not clearly sentence-separated, extract via regex
                    ext = extract_sentences(item)
                    if ext:
                        sentences.extend(ext)
                    else:
                        sentences.append(item)
        if args.sentences_file:
            try:
                with open(args.sentences_file, "r", encoding="utf-8") as fh:
                    content = fh.read()
                sentences.extend(extract_sentences(content))
            except Exception as e:
                logging.error("Failed to read sentences file: %s", e)
                sys.exit(2)
        # Deduplicate and keep order
        seen = set()
        ordered: List[str] = []
        for s in sentences:
            s2 = s.strip()
            if s2 and s2 not in seen:
                ordered.append(s2)
                seen.add(s2)
        if not ordered:
            logging.error("No valid sentences provided. Ensure they end with .!? ")
            sys.exit(2)
        try:
            out_path, total = apply_sentences_docx(args.in_path, args.out_path, ordered)
            logging.info("Sentence mode: converted %d match(es). Wrote: %s", total, out_path)
        except Exception as e:
            logging.error("Sentence processing failed: %s", e)
            sys.exit(1)
        return

    if not args.in_path:
        logging.error("--in is required for CLI mode (omit it when using --gui)")
        sys.exit(2)

    if _zipf_frequency is None:
        logging.error("wordfreq is required for non-sentence CLI mode. Install with: pip install wordfreq")
        sys.exit(2)

    try:
        convert_docx(args.in_path, args.out_path, args.threshold, args.dry_run)
    except FileNotFoundError:
        sys.exit(2)
    except Exception as e:
        logging.error("Processing failed: %s", e)
        sys.exit(1)


if __name__ == "__main__":
    main()
